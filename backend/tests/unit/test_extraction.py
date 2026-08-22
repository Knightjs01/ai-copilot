from io import BytesIO

import pytest
from docx import Document
from fpdf import FPDF

from app.modules.privacy_gateway.exceptions import ExtractionFailedError, UnsupportedFileTypeError
from app.modules.privacy_gateway.extraction import extract_text


def _build_pdf_bytes(text: str) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.cell(text=text)
    return bytes(pdf.output())


def _build_docx_bytes(text: str) -> bytes:
    document = Document()
    document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


_DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def test_extracts_text_from_pdf() -> None:
    content = _build_pdf_bytes("Experienced backend engineer resume content")
    text = extract_text(content=content, content_type="application/pdf")
    assert "Experienced backend engineer resume content" in text


def test_extracts_text_from_docx() -> None:
    content = _build_docx_bytes("Experienced backend engineer resume content")
    text = extract_text(
        content=content,
        content_type=("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
    )
    assert "Experienced backend engineer resume content" in text


def test_unsupported_content_type_raises() -> None:
    with pytest.raises(UnsupportedFileTypeError):
        extract_text(content=b"not a real file", content_type="application/msword")


def test_garbage_pdf_bytes_raise_extraction_failed() -> None:
    with pytest.raises(ExtractionFailedError):
        extract_text(content=b"this is not a valid pdf", content_type="application/pdf")


def test_docx_bullet_list_items_are_marked_as_list_lines() -> None:
    """A Word bullet has no literal bullet glyph in paragraph.text -- it's drawn from list
    formatting -- so plain extraction used to silently flatten it into an indistinguishable
    paragraph. Confirms the fix marks it with a "- " prefix the frontend can render as a real
    list, without inventing a marker on paragraphs that were never a list to begin with."""
    document = Document()
    document.add_paragraph("Responsibilities")
    document.add_paragraph("Own the platform roadmap", style="List Bullet")
    document.add_paragraph("Mentor senior engineers", style="List Bullet")
    document.add_paragraph("A closing paragraph, not a list item.")
    buffer = BytesIO()
    document.save(buffer)

    text = extract_text(content=buffer.getvalue(), content_type=_DOCX_CONTENT_TYPE)
    lines = text.split("\n")

    assert "Responsibilities" in lines
    assert "- Own the platform roadmap" in lines
    assert "- Mentor senior engineers" in lines
    assert "A closing paragraph, not a list item." in lines


def test_docx_blank_paragraphs_still_separate_sections() -> None:
    """Regression guard: the rewrite must keep blank-line spacing between paragraphs, not just
    the new bullet-marking behavior."""
    document = Document()
    document.add_paragraph("First section")
    document.add_paragraph("")
    document.add_paragraph("Second section")
    buffer = BytesIO()
    document.save(buffer)

    text = extract_text(content=buffer.getvalue(), content_type=_DOCX_CONTENT_TYPE)
    assert "First section\n\nSecond section" in text
